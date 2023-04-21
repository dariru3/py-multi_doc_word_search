import docx
from reportlab.pdfgen import canvas

def import_txt(path):
    with open(path, 'r') as txt_file:
        text = txt_file.read()
        return text

def create_sample_pdf(filename, text_file):
    pdf = canvas.Canvas(filename)
    y = 750
    with open(text_file, 'r') as f:
        for line in f:
            pdf.drawString(100, y, line.strip())
            y -= 20
    pdf.save()

def create_sample_docx(filename, text_file):
    # Create a new Word document
    doc = docx.Document()

    # Add a new paragraph to the document
    paragraph = doc.add_paragraph()

    # Add text to the paragraph
    text = import_txt(text_file)
    paragraph.add_run(text)

    # Save the document to disk
    doc.save(filename)

# create_sample_pdf("example1.pdf", "text1.txt")
# create_sample_docx("example2.docx", "text2.txt")